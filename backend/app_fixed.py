from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime
import pandas as pd
import os
from werkzeug.utils import secure_filename
import json

from utils.docx_filler import fill_document


app = Flask(__name__)
CORS(app)

# Database configuration
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///timesheet.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['DOCUMENTS_FOLDER'] = 'documents'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

db = SQLAlchemy(app)

#fill document function
from docx import Document
from copy import deepcopy
from docx import Document
from copy import deepcopy

def fill_document(template_path, output_path, data, client_type, employees=None):
    doc = Document(template_path)

    # Fill paragraphs first
    for paragraph in doc.paragraphs:
        full_text = paragraph.text
        needs_replacement = False
        new_text = full_text
        
        for key, value in data.items():
            if key in new_text:
                new_text = new_text.replace(key, value)
                needs_replacement = True
                print(f"DEBUG paragraph: Replacing '{key}' with '{value}'")
        
        if needs_replacement:
            for run in paragraph.runs:
                run.text = ''
            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)
    
    # Fill table and add employee rows
    for table in doc.tables:
        # First, add employee rows if provided
        if employees:
            add_employee_rows(table, employees, client_type)
        
        # Then fill all remaining placeholders in the table
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Get the full text of the paragraph
                    full_text = paragraph.text
                    
                    # Debug: print cells that contain placeholders
                    if '[' in full_text and ']' in full_text:
                        print(f"DEBUG table cell: Found placeholder: '{full_text}'")
                    
                    # Check if any placeholder exists
                    needs_replacement = False
                    new_text = full_text
                    
                    for key, value in data.items():
                        if key in new_text:
                            print(f"DEBUG table cell: Replacing '{key}' with '{value}'")
                            new_text = new_text.replace(key, value)
                            needs_replacement = True
                    
                    # If replacement is needed, update the paragraph
                    if needs_replacement:
                        # Clear all runs
                        for run in paragraph.runs:
                            run.text = ''
                        
                        # Set the new text in the first run (or create one if none exist)
                        if paragraph.runs:
                            paragraph.runs[0].text = new_text
                        else:
                            paragraph.add_run(new_text)
    
    doc.save(output_path)

def add_employee_rows(table, employees, client_type):
    """
    Add multiple employee rows to the invoice table
    """
    # Find the template row with [name] placeholder
    template_row_index = None
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if "[name]" in cell.text:
                template_row_index = i
                break
        if template_row_index is not None:
            break
    
    if template_row_index is None:
        print("Template row with [name] not found")
        return
    
    template_row = table.rows[template_row_index]
    
    # Fill first employee in the template row
    fill_employee_row(template_row, 1, employees[0], client_type)
    
    # Add additional rows for remaining employees
    for idx in range(1, len(employees)):
        # Add a new row by copying the table row element
        new_row = add_row_after(table, template_row_index + idx - 1)
        fill_employee_row(new_row, idx + 1, employees[idx], client_type)

def add_row_after(table, row_index):
    """Add a new row after the specified index by cloning"""
    # Get the row to clone
    row_to_clone = table.rows[row_index]
    
    # Clone the row element
    tbl = table._tbl
    new_tr = deepcopy(row_to_clone._tr)
    
    # Insert after the specified row
    tbl.insert(row_index + 2, new_tr)
    
    return table.rows[row_index + 1]

def fill_employee_row(row, serial_no, emp, client_type):
    """Fill a row with employee data by replacing placeholders"""
    cells = row.cells
    
    print(f"DEBUG fill_employee_row: Processing row with {len(cells)} cells for client_type={client_type}")
    
    # Data mapping based on client type
    if client_type == 'same_state' or client_type == 'other_state':
        # INR Invoice: S.No, Name, DOJ, Total Days, Working Days, Status, Location, Net Amount
        replacements = {
            0: str(serial_no),
            1: str(emp.get('name', '') or ''),
            2: str(emp.get('date_of_joining', '') or ''),
            3: str(emp.get('total_days', '') or ''),
            4: str(emp.get('working_days', '') or ''),
            5: str(emp.get('status', 'Active') or 'Active'),
            6: str(emp.get('location', '') or ''),
            7: str(emp.get('net_amount', '') or '')
        }
        
        # Replace content in each cell
        for cell_idx, new_value in replacements.items():
            if cell_idx < len(cells):
                cell = cells[cell_idx]
                
                # Clear the cell content and set new value
                for paragraph in cell.paragraphs:
                    if paragraph.runs:
                        first_run = paragraph.runs[0]
                        for run in paragraph.runs:
                            run.text = ''
                        first_run.text = new_value
                    else:
                        paragraph.text = new_value
    else:  # foreign - USD Invoice
        replacements = {
            0:str(emp.get('name', '') or ''),
            1:str(emp.get('total_hours', '') or ''),
            2:str(emp.get('rate_per_hour', '') or '') + 'USD/hr',  # Add /hr suffix
            3:str(emp.get('net_amount', '') or '')
        }
        
        # Replace content in each cell
        for cell_idx, new_value in replacements.items():
            if cell_idx < len(cells):
                cell = cells[cell_idx]
                
                # Clear the cell content and set new value
                for paragraph in cell.paragraphs:
                    if paragraph.runs:
                        first_run = paragraph.runs[0]
                        for run in paragraph.runs:
                            run.text = ''
                        first_run.text = new_value
                    else:
                        paragraph.text = new_value
# Create folders
for folder in [app.config['UPLOAD_FOLDER'], app.config['DOCUMENTS_FOLDER']]:
    if not os.path.exists(folder):
        os.makedirs(folder)

# Database Models
class Company(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(200), nullable=False)
    contact_number = db.Column(db.String(20), nullable=False)
    building_no = db.Column(db.String(50), nullable=False)
    local_street = db.Column(db.String(50), nullable=False)
    city = db.Column(db.String(50), nullable=False)
    state = db.Column(db.String(50), nullable=False)
    country = db.Column(db.String(50), nullable=False)
    GST = db.Column(db.String(50))
    SAC = db.Column(db.String(50))
    email = db.Column(db.String(100), nullable=False)
    client_type = db.Column(db.String(50), nullable=False)  # same_state, other_state, foreign
    document_path = db.Column(db.String(500))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    po_numbers = db.relationship('PONumber', backref='company', lazy=True, cascade='all, delete-orphan')

class PONumber(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    company_id = db.Column(db.Integer, db.ForeignKey('company.id'), nullable=False)
    po_number = db.Column(db.String(50), nullable=False)
    monthly_budget = db.Column(db.Float)  # For Indian clients
    hourly_rate = db.Column(db.Float)  # For foreign clients
    igst = db.Column(db.Float, default=18)
    cgst = db.Column(db.Float, default=9)
    sgst = db.Column(db.Float, default=9)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    employees = db.relationship('Employee', backref='po_number', lazy=True, cascade='all, delete-orphan')
    invoices = db.relationship('Invoice', backref='po_number', lazy=True, cascade='all, delete-orphan')

class Employee(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    po_id = db.Column(db.Integer, db.ForeignKey('po_number.id'), nullable=False)
    name = db.Column(db.String(200), nullable=False)
    email = db.Column(db.String(100))
    location = db.Column(db.String(100), nullable=True)
    date_of_joining = db.Column(db.String(50),nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class Invoice(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    company_id = db.Column(db.Integer, db.ForeignKey('company.id'), nullable=False)
    po_id = db.Column(db.Integer, db.ForeignKey('po_number.id'), nullable=False)
    invoice_number = db.Column(db.String(100), unique=True, nullable=False)
    invoice_data = db.Column(db.Text)  # JSON data
    total_amount = db.Column(db.Float)
    sub_total = db.Column(db.Float)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    month = db.Column(db.String(20))
    year = db.Column(db.Integer)
    company = db.relationship('Company', backref='invoices')

# Create tables
with app.app_context():
    db.create_all()

# Helper functions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'xlsx', 'xls', 'pdf', 'doc', 'docx'}

def parse_hours(hours_value):
    if not hours_value or pd.isna(hours_value):
        return 0
    hours_str = str(hours_value).lower()
    if 'hours' in hours_str or 'hour' in hours_str:
        import re
        match = re.search(r'(\d+(?:\.\d+)?)', hours_str)
        return float(match.group(1)) if match else 0
    try:
        return float(hours_str)
    except (ValueError, TypeError):
        return 0

def process_timesheet(file_path, po_data, client_type):
    df_header = pd.read_excel(file_path, header=None)
    employee_name = df_header.iloc[1, 1] if len(df_header) > 1 else "Unknown Employee"
    
    df = pd.read_excel(file_path, skiprows=4)
    total_days = 22  # Default total days in month
    hours_col = None
    for col in df.columns:
        if 'regular hours' in str(col).lower() or 'hours worked' in str(col).lower():
            hours_col = col
            break
    
    if hours_col is None:
        hours_col = 'Regular hours worked'
    
    df = df.dropna(subset=[hours_col])
    
    result = {'employee_name': employee_name}
    
    if client_type == 'foreign':
        total_hours = sum(parse_hours(hours) for hours in df[hours_col])
        rate = po_data.hourly_rate or 0
        total_amount = total_hours * rate
        
        result.update({
            'total_worked_hours': total_hours,
            'total_amount': total_amount,
            'sub_total': total_amount,
            'toal_days': total_days,
            'calculation_type': 'hourly'
        })
    else:
        total_worked_days = sum(1 for hours in df[hours_col] if hours and str(hours).strip())
        total_budget = po_data.monthly_budget or 0
        per_day_budget = total_budget / 22
        total_amount = per_day_budget * total_worked_days
        
        result.update({
            'total_worked_days': total_worked_days,
            'per_day_budget': per_day_budget,
            'total_amount': total_amount,
            'calculation_type': 'daily'
        })
        
        if client_type == 'same_state':
            igst_amount = total_amount * (po_data.igst / 100)
            sub_total = total_amount + igst_amount
            result.update({'IGST': igst_amount, 'sub_total': sub_total})
        else:
            cgst_amount = total_amount * (po_data.cgst / 100)
            sgst_amount = total_amount * (po_data.sgst / 100)
            sub_total = total_amount + cgst_amount + sgst_amount
            result.update({'CGST': cgst_amount, 'SGST': sgst_amount, 'sub_total': sub_total})
    
    return result

# API Endpoints

@app.route('/api/companies', methods=['POST'])
def create_company():
    try:
        data = request.form
        file = request.files.get('document')
        
        # Validate required fields
        if not data.get('name') or not data.get('contact_number') or not data.get('email'):
            return jsonify({'error': 'Missing required fields'}), 400
        
        company = Company(
            name=data.get('name'),
            contact_number=data.get('contact_number'),
            building_no = data.get('building_no'),
            local_street = data.get('local_street'),
            city = data.get('city'),
            state = data.get('state'),
            country = data.get('country'),
            GST=data.get('GST'),
            SAC=data.get('SAC'),
            email=data.get('email'),
            client_type=data.get('client_type')
        )
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['DOCUMENTS_FOLDER'], f"{datetime.now().timestamp()}_{filename}")
            file.save(filepath)
            company.document_path = filepath
        
        db.session.add(company)
        db.session.flush()
        
        po_numbers_str = data.get('po_numbers', '[]')
        print(f"Received PO numbers: {po_numbers_str}")  # Debug log
        
        po_numbers = json.loads(po_numbers_str)
        for po_data in po_numbers:
            # Convert empty strings to None for numeric fields
            monthly_budget = po_data.get('monthly_budget')
            hourly_rate = po_data.get('hourly_rate')
            
            po = PONumber(
                company_id=company.id,
                po_number=po_data['po_number'],
                monthly_budget=float(monthly_budget) if monthly_budget and monthly_budget != '' else None,
                hourly_rate=float(hourly_rate) if hourly_rate and hourly_rate != '' else None,
                igst=float(po_data.get('igst', 18)),
                cgst=float(po_data.get('cgst', 9)),
                sgst=float(po_data.get('sgst', 9))
            )
            db.session.add(po)
            db.session.flush()
            
            employees = po_data.get('employees', [])
            for emp_data in employees:
                if emp_data.get('name'):  # Only add if name exists
                    employee = Employee(
                        po_id=po.id,
                        name=emp_data['name'],
                        email=emp_data.get('email', ''),
                        date_of_joining=emp_data.get('doj',''),
                        location=emp_data.get('location','')
                    )
                    db.session.add(employee)
        
        db.session.commit()
        return jsonify({'message': 'Company created successfully', 'company_id': company.id}), 201
    
    except json.JSONDecodeError as e:
        db.session.rollback()
        print(f"JSON decode error: {str(e)}")
        return jsonify({'error': f'Invalid JSON format: {str(e)}'}), 400
    except Exception as e:
        db.session.rollback()
        print(f"Error: {str(e)}")  # Debug log
        return jsonify({'error': str(e)}), 400

@app.route('/api/companies', methods=['GET'])
def get_companies():
    companies = Company.query.all()
    return jsonify([{
        'id': c.id,
        'name': c.name,
        'contact_number': c.contact_number,
        'email': c.email,
        'client_type': c.client_type,
        'created_at': c.created_at.isoformat(),
        'po_count': len(c.po_numbers)
    } for c in companies])

@app.route('/api/companies/<int:company_id>', methods=['GET'])
def get_company(company_id):
    company = Company.query.get_or_404(company_id)
    return jsonify({
        'id': company.id,
        'name': company.name,
        'contact_number': company.contact_number,
        'email': company.email,
        'client_type': company.client_type,
        'created_at': company.created_at.isoformat(),
        'po_numbers': [{
            'id': po.id,
            'po_number': po.po_number,
            'monthly_budget': po.monthly_budget,
            'hourly_rate': po.hourly_rate,
            'igst': po.igst,
            'cgst': po.cgst,
            'sgst': po.sgst,
            'employees': [{'id': e.id, 'name': e.name, 'email': e.email} for e in po.employees]
        } for po in company.po_numbers]
    })

@app.route('/api/companies/<int:company_id>/po-numbers', methods=['GET'])
def get_po_numbers(company_id):
    company = Company.query.get_or_404(company_id)
    return jsonify([{
        'id': po.id,
        'po_number': po.po_number,
        'monthly_budget': po.monthly_budget,
        'hourly_rate': po.hourly_rate,
        'employee_count': len(po.employees)
    } for po in company.po_numbers])

@app.route('/api/po-numbers/<int:po_id>/employees', methods=['GET'])
def get_po_employees(po_id):
    po = PONumber.query.get_or_404(po_id)
    return jsonify([{
        'id': e.id,
        'name': e.name,
        'email': e.email,
        'date_of_joining': e.date_of_joining,
        'location': e.location
    } for e in po.employees])

@app.route('/api/invoices/generate', methods=['POST'])
def generate_invoice():
    try:
        company_id = request.form.get('company_id')
        po_id = request.form.get('po_id')
        month = request.form.get('month')
        year = request.form.get('year')
        
        company = Company.query.get_or_404(company_id)
        po = PONumber.query.get_or_404(po_id)
        
        files = request.files.getlist('files')
        results = []
        
        for file in files:
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(filepath)
                
                try:
                    result = process_timesheet(filepath, po, company.client_type)
                    result['filename'] = filename
                    result['filepath'] = filepath  # store full path so download can find it reliably
                    results.append(result)
                    # DO NOT delete the file here — keep it for download/generation
                except Exception as e:
                    results.append({'filename': filename, 'error': str(e)})
        
        grand_total = {
            'total_amount': sum(r.get('total_amount', 0) for r in results if 'error' not in r),
            'sub_total': sum(r.get('sub_total', 0) for r in results if 'error' not in r),
            'total_hours': sum(r.get('total_worked_hours', 0) for r in results if 'error' not in r),
            'total_days': sum(r.get('total_days', 0) for r in results if 'error' not in r),
            'IGST': sum(r.get('IGST', 0) for r in results if 'error' not in r),
            'CGST': sum(r.get('CGST', 0) for r in results if 'error' not in r),
            'SGST': sum(r.get('SGST', 0) for r in results if 'error' not in r)
        }
        
        invoice_number = f"INV-{company.id}-{po_id}-{year}{month}-{datetime.now().strftime('%H%M%S')}"
        
        invoice = Invoice(
            company_id=company_id,
            po_id=po_id,
            invoice_number=invoice_number,
            invoice_data=json.dumps({'employees': results, 'grand_total': grand_total}),
            total_amount=grand_total['total_amount'],
            sub_total=grand_total['sub_total'],
            month=month,
            year=int(year)
        )
        
        db.session.add(invoice)
        db.session.commit()
        
        return jsonify({
            'invoice_id': invoice.id,
            'invoice_number': invoice_number,
            'company': {'name': company.name, 'email': company.email, 'client_type': company.client_type},
            'po_number': po.po_number,
            'employees': results,
            'grand_total': grand_total,
            'month': month,
            'year': year
        })
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 400


@app.route('/api/invoices', methods=['GET'])
def get_invoices():
    invoices = Invoice.query.order_by(Invoice.created_at.desc()).all()
    return jsonify([{
        'id': inv.id,
        'invoice_number': inv.invoice_number,
        'company_name': inv.company.name,
        'po_number': inv.po_number.po_number,
        'total_amount': inv.total_amount,
        'sub_total': inv.sub_total,
        'month': inv.month,
        'year': inv.year,
        'created_at': inv.created_at.isoformat()
    } for inv in invoices])

@app.route('/api/invoices/<int:invoice_id>', methods=['GET'])
def get_invoice(invoice_id):
    invoice = Invoice.query.get_or_404(invoice_id)
    invoice_data = json.loads(invoice.invoice_data)
    
    return jsonify({
        'id': invoice.id,
        'invoice_number': invoice.invoice_number,
        'company': {
            'name': invoice.company.name,
            'email': invoice.company.email,
            'contact_number': invoice.company.contact_number,
            'client_type': invoice.company.client_type
        },
        'po_number': invoice.po_number.po_number,
        'employees': invoice_data['employees'],
        'grand_total': invoice_data['grand_total'],
        'month': invoice.month,
        'year': invoice.year,
        'created_at': invoice.created_at.isoformat()
    })

@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({'status': 'ok'})

from docx import Document
from docx.shared import Inches
@app.route('/api/invoices/<int:invoice_id>/download-docx', methods=['GET'])
def download_invoice_docx(invoice_id):
    try:
        invoice = Invoice.query.get_or_404(invoice_id)
        company = invoice.company
        po = invoice.po_number

        # Fetch rates and taxes from PO table
        hourly_rate = po.hourly_rate or 0
        monthly_budget = po.monthly_budget or 0
        cgst_rate = po.cgst or 9
        sgst_rate = po.sgst or 9
        igst_rate = po.igst or 18

        client_type = company.client_type
        print(f"DEBUG: Client Type = {client_type}")

        # Parse multiple Excel files from invoice_data
        invoice_data = json.loads(invoice.invoice_data or "{}")
        employee_entries = invoice_data.get("employees", [])

        if not employee_entries:
            return jsonify({'error': 'No Excel files linked to this invoice'}), 404

        total_invoice_amount = 0
        total_cgst = total_sgst = total_igst = 0
        total_days = 22
        all_employees = []

        # Fetch employee info from DB
        employees_list = Employee.query.filter_by(po_id=po.id).all()

        # Process each Excel sheet
        for idx, emp_data in enumerate(employee_entries, start=1):
            excel_filename = emp_data.get("filename")
            if not excel_filename:
                continue

            excel_path = os.path.join(app.config['UPLOAD_FOLDER'], excel_filename)
            if not os.path.exists(excel_path):
                print(f"DEBUG: Excel file not found: {excel_path}")
                continue

            # Read Excel
            df_header = pd.read_excel(excel_path, header=None)
            employee_name = df_header.iloc[1, 1] if len(df_header) > 1 else f"Employee {idx}"
            location_name = df_header.iloc[3, 1] if len(df_header) > 2 else ""
            df = pd.read_excel(excel_path, skiprows=4)

            # Detect 'Regular Hours' column
            hours_col = None
            for col in df.columns:
                if 'regular hours' in str(col).lower():
                    hours_col = col
                    break
            if not hours_col:
                hours_col = 'Regular hours worked'

            df = df.dropna(subset=[hours_col])

            total_worked_hours = 0
            total_worked_days = 0
            total_amount = sub_total = 0
            CGST = SGST = IGST = 0

            # Calculate totals
            if client_type == "foreign":
                for hours in df[hours_col]:
                    hours_str = str(hours).lower()
                    if "hours" in hours_str or "hour" in hours_str:
                        # Extract numeric value
                        import re
                        match = re.search(r'(\d+(?:\.\d+)?)', hours_str)
                        if match:
                            total_worked_hours += float(match.group(1))
                
                total_amount = total_worked_hours * hourly_rate
                sub_total = total_amount
                print(f"DEBUG Foreign: Hours={total_worked_hours}, Rate={hourly_rate}, Amount={total_amount}")
            else:
                for hours in df[hours_col]:
                    if "hours" in str(hours):
                        total_worked_days += 1
                per_day_budget = monthly_budget / 22 if monthly_budget else 0
                total_amount = total_worked_days * per_day_budget

                if client_type == "other_state":
                    cgst_amt = total_amount * (cgst_rate / 100)
                    sgst_amt = total_amount * (sgst_rate / 100)
                    sub_total = total_amount + cgst_amt + sgst_amt
                    CGST, SGST = cgst_amt, sgst_amt
                else:
                    igst_amt = total_amount * (igst_rate / 100)
                    sub_total = total_amount + igst_amt
                    IGST = igst_amt

            # Accumulate totals
            total_invoice_amount += total_amount
            total_cgst += CGST
            total_sgst += SGST
            total_igst += IGST

            # Get employee info from DB or use Excel data
            if idx <= len(employees_list):
                emp = employees_list[idx - 1]
                emp_name = employee_name
                doj = emp.date_of_joining
                location = location_name
            else:
                emp_name = employee_name
                doj = ""
                location = location_name

            # Build employee entry - THIS MUST HAPPEN INSIDE THE LOOP
            if client_type == "same_state" or client_type == "other_state":
                all_employees.append({
                    "name": emp_name,
                    "total_days": total_days,
                    "working_days": total_worked_days,
                    "status": "Active",
                    "date_of_joining": doj,
                    "location": location,
                    "net_amount": f"₹{total_amount:,.2f}"
                })
            else:  # foreign
                emp_dict = {
                    "name": emp_name,
                    "total_hours": f"{total_worked_hours:.2f}",
                    "rate_per_hour": f"{hourly_rate:.2f}",  # Keep $ sign for display
                    "net_amount": f"${total_amount:,.2f}"
                }
                print(f"DEBUG: Adding employee: {emp_dict}")
                print(f"DEBUG: hourly_rate from PO: {hourly_rate}, total_hours: {total_worked_hours}, total_amount: {total_amount}")
                all_employees.append(emp_dict)

        print(f"DEBUG: Total employees to fill: {len(all_employees)}")
        print(f"DEBUG: Employee data: {all_employees}")

        # Prepare totals for invoice
        grand_total = total_invoice_amount + total_cgst + total_sgst + total_igst

        # Generate DOCX
        if client_type == "other_state":
            template_path = os.path.join("templates", "INR INVOICE.docx")
            data = {
                "[Invoice number]": invoice.invoice_number,
                "[Date]": invoice.created_at.strftime("%Y-%m-%d"),
                "[MM]": invoice.created_at.strftime("%m"),
                "[YYYY]": invoice.created_at.strftime("%Y"),
                "[PO number]": po.po_number,
                "[company_name]": company.name,
                "[building_no]" : company.building_no,
                "[local_street]" : company.local_street,
                "[city]" : company.city,
                "[state]" : company.state,
                "[country]" : company.country,
                "[GST]": company.GST,
                "[SAC]": company.SAC,
                "[sub_total]": f"₹{total_invoice_amount:,.2f}",  # Try without brackets
                "[CGST]": f"₹{total_cgst:,.2f}",
                "[SGST]": f"₹{total_sgst:,.2f}",
                "[IGST]": f"₹{total_igst:,.2f}",
                "[TIA]": f"₹{grand_total:,.2f}"
            }
        elif client_type == "same_state":
            template_path = os.path.join("templates", "same_state.docx")
            data = {
                "[Invoice number]": invoice.invoice_number,
                "[Date]": invoice.created_at.strftime("%Y-%m-%d"),
                "[MM]": invoice.created_at.strftime("%m"),
                "[YYYY]": invoice.created_at.strftime("%Y"),
                "[PO number]": po.po_number,
                "[company_name]": company.name,
                "[building_no]" : company.building_no,
                "[local_street]" : company.local_street,
                "[city]" : company.city,
                "[state]" : company.state,
                "[country]" : company.country,
                "[GST]": company.GST,
                "[SAC]": company.SAC,
                "[sub_total]": f"₹{total_invoice_amount:,.2f}",  # Try without brackets
                "[CGST]": f"₹{total_cgst:,.2f}",
                "[SGST]": f"₹{total_sgst:,.2f}",
                "[IGST]": f"₹{total_igst:,.2f}",
                "[TIA]": f"₹{grand_total:,.2f}"
            }
        else :
            template_path = os.path.join("templates", "USD INVOICE.docx")
            data = {
                '[Date]': invoice.created_at.strftime("%Y-%m-%d"),
                '[PO number]': str(po.po_number),
                '[ST]': f"${grand_total:,.2f}",
                "[Invoice number]": invoice.invoice_number,
                "[Date]": invoice.created_at.strftime("%Y-%m-%d"),
                "[MM]": invoice.created_at.strftime("%m"),
                "[YYYY]": invoice.created_at.strftime("%Y"),
                "[PO number]": po.po_number,
                "[company_name]": company.name,  # This is the invoice number placeholder # For the payable line
            }
            print(f"DEBUG: Template data: {data}")

        if not os.path.exists(template_path):
            return jsonify({'error': f'Invoice template not found at {template_path}'}), 404

        output_filename = f"Invoice_{invoice.invoice_number}.docx"
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)

        print(f"DEBUG: Calling fill_document with {len(all_employees)} employees")
        fill_document(template_path, output_path, data, client_type, all_employees)
        print(f"DEBUG: Document created at {output_path}")

        return send_file(output_path, as_attachment=True, download_name=output_filename)

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/companies/<int:company_id>', methods=['DELETE'])
def delete_company(company_id):
    try:
        company = Company.query.get_or_404(company_id)

        # Collect related invoices before deletion to clean files
        invoices = Invoice.query.filter_by(company_id=company_id).all()

        # Attempt to remove generated DOCX files and any stored timesheet files
        for inv in invoices:
            try:
                output_filename = f"Invoice_{inv.invoice_number}.docx"
                output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
                if os.path.exists(output_path):
                    os.remove(output_path)
            except Exception:
                pass

            try:
                data = json.loads(inv.invoice_data or '{}')
                for entry in data.get('employees', []):
                    fp = entry.get('filepath')
                    if fp and os.path.exists(fp):
                        os.remove(fp)
            except Exception:
                pass

        # Remove company's uploaded document if present
        if company.document_path and os.path.exists(company.document_path):
            try:
                os.remove(company.document_path)
            except Exception:
                pass

        # Delete invoices explicitly
        Invoice.query.filter_by(company_id=company_id).delete()

        # Delete company (cascades to POs and employees)
        db.session.delete(company)
        db.session.commit()
        return jsonify({'message': 'Company and related data deleted'}), 200
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)